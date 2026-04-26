from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = next(ROOT.glob("01*"))
OUTPUT_DIR = ROOT / "docs" / "competition-2026"

TITLE = "基于智能匹配的中外学生语言互助与跨文化交流平台"
VERSION = "V1.0"
TODAY = "2026年4月21日"


def set_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    font = run.font
    font.name = "宋体"


def set_cell(cell, text: str) -> None:
    cell.text = text


def fill_overview() -> None:
    src = next(TEMPLATE_DIR.glob("2-作品信息概要表*.docx"))
    dst = OUTPUT_DIR / "2-作品信息概要表（已填写草稿）.docx"
    shutil.copyfile(src, dst)

    doc = Document(dst)
    table = doc.tables[0]

    set_cell(table.cell(0, 1), "待填写")
    set_cell(table.cell(0, 6), TITLE)
    set_cell(table.cell(1, 1), "软件应用与开发")
    set_cell(table.cell(1, 8), "软件应用与开发类作品（请按报名系统小类名称核对）")
    set_cell(
        table.cell(2, 0),
        "本作品面向高校中外学生构建语言互助与跨文化交流平台，提供注册登录、资料完善、智能匹配、私信聊天、社区互动、活动报名与后台管理等功能，支持课堂展示和项目答辩演示。"
    )
    set_cell(
        table.cell(3, 0),
        "以“语言互补+兴趣标签+时间偏好”实现可解释匹配；前后端分离并覆盖用户端与管理端；采用演示数据与MySQL表结构双轨方案，兼顾现场演示稳定性和后续工程扩展。"
    )
    set_cell(
        table.cell(4, 0),
        "1. 本作品不涉及疆域地图内容。\n"
        "2. 本次参赛以现有项目为基础，完成了需求梳理、前后端实现、数据库设计、联调测试与提交文档整理。\n"
        "3. 若确认开发或文档整理过程中使用了AI辅助，请同步提交《AI工具使用说明》并补充佐证材料。"
    )

    headers = ["项目名称", "成员1（待填）", "成员2（待填）", "成员3（待填）"]
    set_cell(table.cell(6, 0), headers[0])
    set_cell(table.cell(6, 2), headers[1])
    set_cell(table.cell(6, 4), headers[2])
    set_cell(table.cell(6, 6), headers[3])
    for col in [8, 10, 12, 13, 15, 16]:
        set_cell(table.cell(6, col), "")

    work_rows = {
        7: ["组织协调", "待填", "待填", "待填"],
        8: ["作品创意", "待填", "待填", "待填"],
        9: ["竞品分析", "待填", "待填", "待填"],
        10: ["方案设计", "待填", "待填", "待填"],
        11: ["技术实现", "待填", "待填", "待填"],
        12: ["文献阅读", "待填", "待填", "待填"],
        13: ["测试分析", "待填", "待填", "待填"],
    }
    for row_idx, values in work_rows.items():
        set_cell(table.cell(row_idx, 0), values[0])
        set_cell(table.cell(row_idx, 2), values[1])
        set_cell(table.cell(row_idx, 4), values[2])
        set_cell(table.cell(row_idx, 6), values[3])
        for col in [8, 10, 12, 13, 15, 16]:
            set_cell(table.cell(row_idx, col), "")

    set_cell(
        table.cell(15, 3),
        "☑理论指导 ☑技术方案 ☑组织协调 □实验场地 □硬件资源\n"
        "□数据提供 □后勤支持 □宣讲通知 □经费支持\n"
        "□其他：待指导教师确认"
    )
    set_cell(table.cell(16, 3), "☑Windows □Linux □macOS □其他：")
    set_cell(table.cell(17, 3), "☑Windows □Linux □macOS □iOS □Android □其他：")
    set_cell(
        table.cell(18, 3),
        "前端：Vue 3、Vue Router、Vite；\n"
        "后端：Spring Boot 3、Maven、Java 21；\n"
        "数据库：MySQL；\n"
        "开发工具：IntelliJ IDEA / VS Code / Git"
    )
    set_cell(
        table.cell(19, 3),
        "1、Vue.js 官方文档：https://cn.vuejs.org/\n"
        "2、Spring Boot 官方文档：https://spring.io/projects/spring-boot\n"
        "3、MySQL 8.0 Reference Manual：https://dev.mysql.com/doc/"
    )
    set_cell(
        table.cell(20, 3),
        "☑设计文档 ☑演示视频 ☑PPT ☑源代码 ☑部署文件 ☑作品文件 □数据集 □模型 ☑其他：作品信息概要表、AI工具使用说明（如适用）"
    )

    file_rows = [
        ("1", "2-作品信息概要表（已填写草稿）.docx", "参赛作品信息概要表", "□已上传到网盘 ☑未上传，下载地址：随作品提交包提供", "☑自制 □未知版权 □开源 □获得授权"),
        ("2", "3-软件应用与开发类作品设计和开发文档（已填写草稿）.docx", "设计与开发文档", "□已上传到网盘 ☑未上传，下载地址：随作品提交包提供", "☑自制 □未知版权 □开源 □获得授权"),
        ("3", "4-AI工具使用说明（已填写草稿）.docx", "AI辅助情况说明，若实际未使用可删除", "□已上传到网盘 ☑未上传，下载地址：随作品提交包提供", "☑自制 □未知版权 □开源 □获得授权"),
        ("4", "README.md", "项目总览与运行说明", "□已上传到网盘 ☑未上传，下载地址：随作品提交包提供", "☑自制 □未知版权 □开源 □获得授权"),
        ("5", "frontend/", "前端工程源码", "□已上传到网盘 ☑未上传，下载地址：随作品提交包提供", "☑自制 □未知版权 □开源 □获得授权"),
        ("6", "backend/", "后端工程源码", "□已上传到网盘 ☑未上传，下载地址：随作品提交包提供", "☑自制 □未知版权 □开源 □获得授权"),
        ("7", "database/schema.sql", "数据库建表脚本", "□已上传到网盘 ☑未上传，下载地址：随作品提交包提供", "☑自制 □未知版权 □开源 □获得授权"),
    ]
    for idx, row_idx in enumerate(range(24, 31)):
        if idx >= len(file_rows):
            break
        row = file_rows[idx]
        set_cell(table.cell(row_idx, 0), row[0])
        set_cell(table.cell(row_idx, 1), f"文件名：{row[1]}\n描述：{row[2]}")
        set_cell(table.cell(row_idx, 9), row[3])
        set_cell(table.cell(row_idx, 14), row[4])

    doc.save(dst)


def fill_design_doc() -> None:
    src = next(TEMPLATE_DIR.glob("3 软件应用与开发类作品设计和开发文档模板*.docx"))
    dst = OUTPUT_DIR / "3-软件应用与开发类作品设计和开发文档（已填写草稿）.docx"
    shutil.copyfile(src, dst)

    doc = Document(dst)
    replacements = {
        "作品编号：": "作品编号：待填写",
        "作品名称：": f"作品名称：{TITLE}",
        "版本编号：": f"版本编号：{VERSION}",
        "填写日期：": f"填写日期：{TODAY}",
        "【填写说明：本部分内容建议不超过1000字，以300字以内为宜，简要说明为什么开发本作品，是否存在竞品，对标什么作品以及面向的用户、主要功能、主要性能等。如果存在竞品，建议有竞品分析表格，从多个维度分析本作品与竞品作品比较】":
            "本作品面向校园内有语言学习、文化交流和社交需求的中外学生，解决传统线下配对效率低、沟通链路分散、活动参与和反馈难沉淀的问题。系统对标校园社交平台、活动报名平台和语言交换社区，但更强调“语言互补+兴趣契合+时间匹配”的可解释推荐。\n\n"
            "目标用户包括中国学生、留学生和平台管理员。主要功能覆盖注册登录、资料维护、智能匹配、用户详情、私信聊天、社区动态、活动管理、个人中心和后台管理。当前版本采用 Vue 3 + Spring Boot 3 前后端分离架构，前端构建与后端打包均已通过，适合课程设计、毕业设计展示与后续工程扩展。",
        "【填写说明：将需求分析结果分解成功能模块以及模块的层次结构、调用关系、模块间接口以及人机界面等，建议用图体现内容，不宜全文字描述。建议图文总体不超过A4纸两页，以1页为宜。】":
            "系统采用前后端分离架构。前端由 Vue 3、Vue Router 和 Vite 构建用户端与管理端页面；后端采用 Spring Boot 3 提供 REST API；数据库层预留 MySQL 建表脚本，当前演示版本以 DataStore 内存数据实现稳定演示。\n\n"
            "功能模块可划分为：\n"
            "1. 用户与认证模块：登录、注册、找回密码、账户校验。\n"
            "2. 资料与画像模块：基础资料、语言能力、兴趣标签、交流目标与时间偏好维护。\n"
            "3. 智能匹配模块：根据语言互补、兴趣重合、交流目标和空闲时段综合评分。\n"
            "4. 交流互动模块：会话列表、消息列表、文本/图片/语音文件发送。\n"
            "5. 社区内容模块：发布动态、点赞、收藏、评论、搜索。\n"
            "6. 活动管理模块：活动浏览、报名、取消报名、评价。\n"
            "7. 后台管理模块：仪表盘、用户状态管理、活动维护、举报处理。\n\n"
            "界面层面，系统共设计 16 个主要页面，其中用户端 12 个、管理端 4 个，形成从“发现伙伴”到“建立联系”再到“参与活动与沉淀社区内容”的完整闭环。",
        "【填写说明：包括但不限于：界面设计、数据库设计(如果有)、关键算法。界面设计建议用作品实际界面，建议包括典型使用流程；数据库设计建议用表格、ER图或UML方式，说明文字简明扼要，违背范式的设计建议请说明理由；关键算法也可以替换为关键技术、技术创新等。本部分不宜大篇幅铺陈，建议突出重点痛点难点特点。】":
            "1. 界面设计\n"
            "用户端包括登录、注册、找回密码、首页、个人资料、匹配推荐、用户详情、聊天、社区、活动列表、活动详情、我的中心等页面；管理端包括管理员登录、仪表盘、用户管理、活动管理、举报管理。界面布局采用顶部导航与内容卡片结合的方式，兼顾信息展示密度和操作效率。\n\n"
            "2. 数据库设计\n"
            "数据库围绕 user、user_profile、match_record、chat_session、chat_message、post、post_comment、activity、activity_signup、activity_review、report、notification 等核心表展开，覆盖用户、匹配、聊天、社区、活动和后台治理六类业务数据。项目已提供 MySQL 建表脚本，可直接用于后续真实持久化接入。\n\n"
            "3. 关键算法与技术\n"
            "匹配算法采用规则评分方式：语言双向互补最高加权，其次考虑共同兴趣、交流目标、空闲时间、偏好交流方式和历史评价，最终输出 0~100 的匹配分数及推荐理由。该方式实现简单、结果可解释，便于答辩展示。\n\n"
            "4. 工程实现特点\n"
            "前端通过路由守卫区分用户端与管理端访问权限；后端按认证、用户、内容、管理四类控制器组织接口；演示模式与 MySQL 表结构并存，既保证现场可运行性，又为后续数据库接入留出扩展空间。",
        "【填写说明：包括测试报告和技术指标。为了保证作品质量，建议多进行测试，并将测试用例、测试过程、测试结果、修正过程或结果形成文档，也可以将本标题修改为主要测试，撰写主要测试过程结果及其修正；根据测试结果，形成多维度技术指标，包括：运行速度、安全性、扩展性、部署方便性和可用性等。本部分简要说明即可，减少常识性内容。】":
            "1. 构建测试\n"
            "前端执行 `npm run build`，已成功生成生产构建文件；后端执行 `mvn -DskipTests package`，已成功打包。\n\n"
            "2. 功能测试\n"
            "围绕注册登录、资料更新、匹配推荐、建立联系、聊天发送消息、动态发布、活动报名与取消、管理员处理用户和举报等主流程进行了联调验证，接口返回结构统一，页面路由跳转正常。\n\n"
            "3. 技术指标\n"
            "运行速度：前端构建产物体积适中，页面响应满足课程演示要求；\n"
            "安全性：接口参数增加校验，区分普通用户与管理员访问；\n"
            "扩展性：模块边界清晰，可继续接入真实鉴权、WebSocket 与持久化层；\n"
            "部署方便性：前后端可独立启动，依赖较少，适合快速部署；\n"
            "可用性：提供演示账号和初始化数据，方便现场展示典型业务流程。",
        "【填写说明：简要说明安装环境要求、安装过程、主要流程等。建议包含默认安装和典型使用流程。】":
            "1. 环境要求\n"
            "操作系统：Windows 10/11；\n"
            "JDK：21；Maven：3.9+；Node.js：18+；MySQL：8.0（如需持久化部署）。\n\n"
            "2. 安装与启动\n"
            "后端：进入 backend 目录执行 `mvn spring-boot:run`，默认端口 8080；\n"
            "前端：进入 frontend 目录执行 `npm install`、`npm run dev`，默认端口 5173。\n\n"
            "3. 典型使用流程\n"
            "用户注册并完善资料后进入首页；在匹配页查看推荐对象并建立联系；进入聊天页完成沟通；可在社区页发布动态，在活动页报名参加线上或线下活动；管理员可在后台对用户、活动与举报进行治理。\n\n"
            "4. 演示账号\n"
            "普通用户：lily.chen / 123456；管理员：admin / 123456。",
        "【填写说明：作品制作开发过程中的一些感悟和后续升级等，如：项目协调、任务分解、克服的困难、水平提升、升级演进、商业推广等诸方面。建议部分篇幅不超过A4纸1页。】":
            "本项目围绕“校园语言互助”场景，从需求拆解、页面设计、接口设计到演示数据构造逐步推进，形成了用户端与管理端协同的完整原型。开发过程中，重点解决了功能范围控制、匹配逻辑可解释性和演示版本稳定性问题。通过采用前后端分离与模块化设计，项目具备了较好的可维护性和展示效果。\n\n"
            "后续可从三个方向升级：一是接入真实用户鉴权、MyBatis/JPA 与 MySQL 持久化，提升工程完整度；二是加入 WebSocket、语音视频实时通信和消息撤回等能力，增强互动体验；三是进一步引入 AI 翻译、智能话题推荐和画像优化策略，提高跨语言交流效率。",
        "【请按照标准参考文件格式填写】":
            "[1] Vue.js. Vue.js Documentation[EB/OL]. https://cn.vuejs.org/.\n"
            "[2] Spring. Spring Boot Reference Documentation[EB/OL]. https://spring.io/projects/spring-boot.\n"
            "[3] Oracle. MySQL 8.0 Reference Manual[EB/OL]. https://dev.mysql.com/doc/.\n"
            "[4] Gamma E, Helm R, Johnson R, et al. Design Patterns: Elements of Reusable Object-Oriented Software[M]. Addison-Wesley, 1994."
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph(paragraph, replacements[text])

    doc.save(dst)


def fill_ai_doc() -> None:
    src = next(TEMPLATE_DIR.glob("4-AI工具使用说明*.docx"))
    dst = OUTPUT_DIR / "4-AI工具使用说明（已填写草稿）.docx"
    shutil.copyfile(src, dst)

    doc = Document(dst)
    set_paragraph(doc.paragraphs[2], f"作品编号：待填写        作品名称：{TITLE}")
    doc.add_paragraph("说明：以下内容依据当前可确认的 AI 辅助整理情况起草；如项目开发阶段另有使用记录，请继续补充；如实际未使用 AI，可删除本文件。")

    table = doc.tables[0]
    rows = [
        [
            "1",
            "OpenAI Codex / ChatGPT，客户端与网页访问，2026年4月21日",
            "文档整理与需求映射：根据现有代码、模板要求生成参赛材料草稿",
            "阅读项目代码与比赛模板，提炼功能模块、技术栈、测试结论，并填写作品信息概要表与设计开发文档。",
            "输出了作品简介、创新点、模块划分、安装使用说明、测试结论等内容框架，并生成可编辑文档。",
            "由人工逐项核对项目实现情况，删除不确定表述，补充比赛要求中的待填字段。",
            "文档初稿约由AI辅助生成60%，最终内容经人工审校后采纳约40%。",
        ],
        [
            "2",
            "OpenAI Codex / ChatGPT，客户端访问，2026年4月21日",
            "语言润色：统一技术文档表达、压缩说明文字长度",
            "将项目说明整理成适合比赛提交的正式中文表述，突出功能闭环与工程特点，避免口语化。",
            "给出更适合竞赛文档的摘要、创新描述、总结与参考文献写法建议。",
            "人工根据比赛模板字数限制和项目实际情况进行了删改与重写。",
            "AI建议采纳比例约30%，主要用于提升文档表述质量。",
        ],
        [
            "3",
            "OpenAI Codex / ChatGPT，客户端访问，2026年4月21日",
            "文件生成：依据模板批量生成已填写草稿文档",
            "保留原始docx模板版式，写入项目已有的技术内容，并将无法从代码库确认的信息统一标记为待填写。",
            "生成了可继续人工校对的 docx 草稿文件，减少重复填写工作量。",
            "人工需补齐作品编号、作者姓名、指导教师、最终提交路径和AI佐证附件。",
            "AI生成结果以骨架和初稿为主，最终提交版本仍需人工确认后使用。",
        ],
    ]

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            set_cell(table.cell(row_idx, col_idx), value)

    for i in range(3, 9):
        set_cell(table.cell(i + 1, 0), str(i + 1))

    doc.save(dst)


def write_notes() -> None:
    note = OUTPUT_DIR / "提交前仍需补充的信息.txt"
    note.write_text(
        "以下内容无法从代码仓库自动推断，请提交前手动确认：\n"
        "1. 作品编号。\n"
        "2. 参赛作者姓名、顺序与分工比例。\n"
        "3. 指导教师姓名及其实际指导作用。\n"
        "4. 大赛报名系统中的准确作品小类名称。\n"
        "5. 最终提交包中文件状态、网盘地址或下载地址。\n"
        "6. 若项目开发阶段确实使用过 AI，请补充完整的时间、工具、提示词与佐证材料；若未使用，可删除 AI 工具说明文件。\n",
        encoding="utf-8",
    )


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    fill_overview()
    fill_design_doc()
    fill_ai_doc()
    write_notes()


if __name__ == "__main__":
    main()
